"""
Multi-format file processor for VectorDataBuild.

This module handles extraction of text content from various file formats:
- DOCX files (Microsoft Word documents) - extracts paragraphs, headings, tables
- TXT files (plain text) - direct text extraction
- MD files (Markdown) - preserves markdown formatting
- HTML/HTM files (web pages) - converts to markdown, removes scripts/styles
- C++/C/HPP/H files (source code) - formats with syntax highlighting
- XML/JSON files (structured data) - extracts content with structure

The processor extracts text content while preserving structure and formatting,
making it suitable for RAG (Retrieval-Augmented Generation) systems. It supports
batch processing of directories and maintains relative file structure in output.
"""

import json
import os
import sys
from typing import List, Dict, Any, Optional, Tuple
from docx import Document
from loguru import logger
from pathlib import Path
from bs4 import BeautifulSoup
from loguru import logger

from markdownify import markdownify as md

import sys
sys.path.append(os.path.join(os.path.dirname(__file__), '..'))

from utils.config import get_config


class MultiFormatProcessor:
    """Handles extraction of text content from DOCX, TXT, and MD files."""
    
    def __init__(self, language: str = None):
        self.logger = logger.bind(name="MultiFormatProcessor")
        self.supported_extensions = {'.docx', '.txt', '.md','.html', '.htm', '.cpp', '.cc', '.c', '.hpp', '.h', '.xml', '.json', '.qbk'}
        self.language = language
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """
        Extract text content from a DOCX file.
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            Extracted text content
        """
        try:
            doc = Document(docx_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            self.logger.error(f"Error extracting text from {docx_path}: {e}")
            return ""
    
    def extract_text_from_html(self, content: str) -> str:
        """
        Extract text content from a TXT file.
        
        Args:
            txt_path: Path to TXT file
            
        Returns:
            Extracted text content
        """
        try:
            """Convert HTML to Markdown (strip head/script/style/nav/footer)."""
            soup = BeautifulSoup(content, "html.parser")
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            clean_html = str(soup)
            return md(clean_html)
            
        except Exception as e:
            return ""
    
    def extract_text_from_code(self, content: str, extension: str, path: str) -> str:
        """
        Extract text content from a MD file.
        
        Args:
            md_path: Path to MD file
            
        Returns:
            Extracted text content
        """
        try:
            return f"# File: {path}\n\n```{extension.replace(".", "")}\n{content}\n```\n"
            
        except Exception as e:
            return ""
    
    def extract_text_from_qbk(self, qbk_path: str) -> str:
        """
        Extract text content from a QBK (QuickBook) file.
        QBK files are Boost documentation files that contain structured documentation.
        
        Args:
            qbk_path: Path to QBK file
            
        Returns:
            Extracted text content with preserved structure
        """
        try:
            with open(qbk_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # QBK files contain structured documentation with specific syntax
            # Parse common QBK elements and convert to readable format
            processed_content = self._parse_qbk_content(content, qbk_path)
            
            return processed_content
            
        except Exception as e:
            self.logger.error(f"Error extracting text from QBK file {qbk_path}: {e}")
            return ""
    
    def _parse_qbk_content(self, content: str, file_path: str) -> str:
        """
        Parse QBK content and convert to readable markdown format.
        
        Args:
            content: Raw QBK content
            file_path: Path to the QBK file
            
        Returns:
            Processed content in markdown format
        """
        try:
            lines = content.split('\n')
            processed_lines = []
            current_section = ""
            in_code_block = False
            code_block_content = []
            
            # Extract file name for header
            file_name = Path(file_path).stem
            processed_lines.append(f"# {file_name}")
            processed_lines.append("")
            processed_lines.append("*QuickBook Documentation File*")
            processed_lines.append("")
            
            for line in lines:
                line = line.strip()
                
                # Skip empty lines at the beginning
                if not line and not processed_lines:
                    continue
                
                # Handle QBK-specific syntax
                if line.startswith('[') and line.endswith(']'):
                    # Section headers in QBK format
                    section_name = line[1:-1].strip()
                    if section_name:
                        processed_lines.append(f"## {section_name}")
                        processed_lines.append("")
                        current_section = section_name
                
                elif line.startswith('[') and ']' in line:
                    # Inline references or links
                    processed_lines.append(f"**Reference:** {line}")
                    processed_lines.append("")
                
                elif line.startswith('```') or line.startswith('`'):
                    # Code blocks
                    if in_code_block:
                        # End of code block
                        if code_block_content:
                            processed_lines.append("```cpp")
                            processed_lines.extend(code_block_content)
                            processed_lines.append("```")
                            processed_lines.append("")
                        code_block_content = []
                        in_code_block = False
                    else:
                        # Start of code block
                        in_code_block = True
                
                elif in_code_block:
                    # Inside code block
                    code_block_content.append(line)
                
                elif line.startswith('*') and line.endswith('*'):
                    # Emphasis or italic text
                    processed_lines.append(f"*{line[1:-1]}*")
                    processed_lines.append("")
                
                elif line.startswith('**') and line.endswith('**'):
                    # Bold text
                    processed_lines.append(f"**{line[2:-2]}**")
                    processed_lines.append("")
                
                elif line.startswith('#'):
                    # Headers (preserve as-is)
                    processed_lines.append(line)
                    processed_lines.append("")
                
                elif line.startswith('-') or line.startswith('*'):
                    # Lists
                    processed_lines.append(line)
                
                elif line.startswith('>'):
                    # Blockquotes
                    processed_lines.append(line)
                
                elif line and not line.startswith('<!--'):
                    # Regular content (skip HTML comments)
                    # Clean up any remaining QBK-specific syntax
                    clean_line = self._clean_qbk_line(line)
                    if clean_line:
                        processed_lines.append(clean_line)
                        processed_lines.append("")
            
            # Handle any remaining code block
            if in_code_block and code_block_content:
                processed_lines.append("```cpp")
                processed_lines.extend(code_block_content)
                processed_lines.append("```")
                processed_lines.append("")
            
            return '\n'.join(processed_lines)
            
        except Exception as e:
            self.logger.error(f"Error parsing QBK content: {e}")
            # Fallback: return raw content with basic formatting
            return f"# {Path(file_path).stem}\n\n```qbk\n{content}\n```"
    
    def _clean_qbk_line(self, line: str) -> str:
        """
        Clean QBK-specific syntax from a line.
        
        Args:
            line: Raw line from QBK file
            
        Returns:
            Cleaned line suitable for markdown
        """
        try:
            # Remove QBK-specific markers
            line = line.replace('\\n', '\n')
            line = line.replace('\\t', '\t')
            
            # Handle inline code markers
            if '`' in line:
                # Preserve inline code
                pass
            
            # Handle links and references
            if '[[' in line and ']]' in line:
                # Convert QBK link format to markdown
                import re
                line = re.sub(r'\[\[([^\]]+)\]\]', r'[\1](\1)', line)
            
            # Handle emphasis
            if '//' in line and not line.strip().startswith('//'):
                # Convert //emphasis// to *emphasis*
                import re
                line = re.sub(r'//([^/]+)//', r'*\1*', line)
            
            return line.strip()
            
        except Exception as e:
            self.logger.warning(f"Error cleaning QBK line: {e}")
            return line
    
    def extract_text_from_file(self, file_path: str) -> str:
        """
        Extract text content from any supported file type.
        
        Args:
            file_path: Path to file (DOCX, TXT, or MD)
            
        Returns:
            Extracted text content
        """
        file_path_obj = Path(file_path)
        extension = file_path_obj.suffix.lower()
        
        if extension == '.docx':
            return self.extract_text_from_docx(file_path)
        elif extension == '.qbk':
            return self.extract_text_from_qbk(file_path)
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read().strip()
            
        except Exception as e:
            self.logger.error(f"Error extracting text from {file_path}: {e}")
            return ""
        
        if extension == '.txt':
            return content
        elif extension == '.md':
            return content
        elif extension in [".html", ".htm"]:
            return self.extract_text_from_html(content)
        elif extension in [".cpp", ".cc", ".c", ".h",".json",".xml", ".hpp"]:
            return content
        else:
            self.logger.warning(f"Unsupported file type: {extension}")
            return ""
    
    def process_file(self, file_path: str, output_dir: str) -> bool:
        """
        Process a single file (DOCX, TXT, or MD) and save as markdown.
        
        Args:
            file_path: Path to file
            output_dir: Directory to save processed file
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Extract text
            text_content = self.extract_text_from_file(file_path)
            
            if not text_content:
                self.logger.warning(f"No text content extracted from {file_path}")
                return False
            
            # Create output path
            input_file = Path(file_path)
            output_path = Path(output_dir)
            output_path.mkdir(parents=True, exist_ok=True)
            
            # Save as markdown
            md_file = output_path / f"{input_file.stem}.md"
            with open(md_file, 'w', encoding='utf-8') as f:
                f.write(f"# {input_file.stem}\n\n{text_content}")
            
            self.logger.info(f"Processed {file_path} -> {md_file}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error processing {file_path}: {e}")
            return False
    
    def get_file_list(self, input_path: str, max_files: Optional[int] = None) -> List[str]:
        """
        Process all supported files (DOCX, TXT, MD) in a directory.
        
        Args:
            input_dir: Directory containing files
            output_dir: Directory to save processed files
            max_files: Maximum number of files to process (None for all)
            
        Returns:
            List of file paths
        """
        try:
            
            # Find all supported files
            all_files = []
            for extension in self.supported_extensions:
                all_files.extend(list(input_path.rglob(f"*{extension}")))
            
            if max_files:
                all_files = all_files[:max_files]
            
            self.logger.info(f"Found {len(all_files)} files to process")
            
            return all_files
            
            
            processed_count = 0
            failed_count = 0
            total_size = 0
            file_type_counts = {'.docx': 0, '.txt': 0, '.md': 0,'.html': 0, '.htm': 0, '.cpp': 0, '.cc': 0, '.c': 0, '.hpp': 0, '.h': 0, '.xml': 0, '.json': 0, '.qbk': 0}
            
            for i, file_path in enumerate(all_files, 1):
                try:
                    self.logger.info(f"Processing {i}/{len(all_files)}: {file_path.name}")
                    
                    # Extract text
                    text_content = self.extract_text_from_file(str(file_path))
                    
                    if text_content:
                        # Save as markdown
                        relative_path = file_path.relative_to(self.input_path)
                        md_path = output_path / relative_path.with_suffix('.md')
                        md_path.parent.mkdir(parents=True, exist_ok=True)
                        
                        with open(md_path, 'w', encoding='utf-8') as f:
                            f.write(f"# {file_path.stem}\n\n{text_content}")
                        
                        processed_count += 1
                        total_size += len(text_content)
                        file_type_counts[file_path.suffix.lower()] += 1
                        output_file_list.append(md_path)
                    else:
                        failed_count += 1
                        self.logger.warning(f"No content extracted from {file_path}")
                        
                except Exception as e:
                    failed_count += 1
                    self.logger.error(f"Error processing {file_path}: {e}")
                    continue
            
            stats = {
                'total_files': len(all_files),
                'processed': processed_count,
                'failed': failed_count,
                'total_text_size': total_size,
                'file_type_counts': file_type_counts,
                'output_directory': str(output_path)
            }
            
            self.logger.info(f"Processing complete: {processed_count} processed, {failed_count} failed")
            self.logger.info(f"File type breakdown: {file_type_counts}")
            return stats, output_file_list
            
        except Exception as e:
            self.logger.error(f"Error processing directory: {e}")
            return {'error': str(e)}
    
    def create_processing_summary(self, stats: Dict[str, Any], output_dir: str) -> bool:
        """
        Create a summary of the processing results.
        
        Args:
            stats: Processing statistics
            output_dir: Directory to save summary
            
        Returns:
            True if successful, False otherwise
        """
        try:
            output_path = Path(output_dir)
            summary_file = output_path / "processing_summary.json"
            
            with open(summary_file, 'w', encoding='utf-8') as f:
                json.dump(stats, f, ensure_ascii=False, indent=2)
            
            self.logger.info(f"Processing summary saved to {summary_file}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error creating summary: {e}")
            return False


def main():
    """Main function for command-line usage."""
    import argparse
    
    parser = argparse.ArgumentParser(description="Process DOCX, TXT, and MD files to markdown")
    parser.add_argument("--input-dir", default = "scraped_files", required=True, help="Input directory containing files")
    parser.add_argument("--output-dir", required=True, help="Output directory for markdown files")
    parser.add_argument("--max-files", type=int, help="Maximum number of files to process")
    
    args = parser.parse_args()
    
    processor = MultiFormatProcessor()
    stats = processor.process_directory(args.input_dir, args.output_dir, args.max_files)
    
    if 'error' in stats:
        print(f"Processing failed: {stats['error']}")
        exit(1)
    
    # Create summary
    processor.create_processing_summary(stats, args.output_dir)
    
    print(f"Processing completed:")
    print(f"  Total files: {stats['total_files']}")
    print(f"  Processed: {stats['processed']}")
    print(f"  Failed: {stats['failed']}")
    print(f"  Total text size: {stats['total_text_size']} characters")
    print(f"  File type breakdown: {stats.get('file_type_counts', {})}")
    print(f"  Output directory: {stats['output_directory']}")


if __name__ == "__main__":
    main()
